#TODO: Create a letter using starting_letter.txt 
#for each name in invited_names.txt
#Replace the [name] placeholder with the actual name.
#Save the letters in the folder "ReadyToSend".
    
#Hint1: This method will help you: https://www.w3schools.com/python/ref_file_readlines.asp
    #Hint2: This method will also help you: https://www.w3schools.com/python/ref_string_replace.asp
        #Hint3: THis method will help you: https://www.w3schools.com/python/ref_string_strip.asp

from pathlib import Path
from docx import Document

base_path = Path(__file__).parent

invited_names_file_path = base_path/"Input"/"Names"/"invited_names.txt"
starting_letter_file_path = base_path/"Input"/"Letters"/"starting_letter.txt"
output_folder_path = base_path/"Output"/"ReadyToSend"
PLACEHOLDER = "[name]"
print(output_folder_path)

with open(f"{invited_names_file_path}") as names_file:
    names = names_file.readlines()

with open(f"{starting_letter_file_path}") as letter_file:
    letter_data = letter_file.read()
    for name in names:
         replaced_letter_data = letter_data.replace(PLACEHOLDER, f"{name.strip()}")
         with open(f"{output_folder_path}/{name.strip()}.txt","w") as output_file:
            output_file.write(f"{replaced_letter_data}")

from docx import Document
from pathlib import Path

base_path = Path(__file__).parent
invited_names_file_path = base_path / "Input" / "Names" / "invited_names.txt"
starting_letter_file_doc_path = base_path / "Input" / "Letters" / "starting_letter.docx"
output_folder_path = base_path / "Output" / "ReadyToSend"
PLACEHOLDER = "[name]"

with open(invited_names_file_path, encoding="utf-8") as names_file:
    names = names_file.readlines()

# Load the template document
template_doc = Document(starting_letter_file_doc_path)

for name in names:
    name = name.strip()
    new_doc = Document()

    for para in template_doc.paragraphs:
        new_para = para.text.replace(PLACEHOLDER, name)
        new_doc.add_paragraph(new_para)

    new_doc.save(output_folder_path / f"{name}.docx")
